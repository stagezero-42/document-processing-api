from docx import Document
from docx.document import Document as DocObject  # For type hinting
from docx.table import Table as TableObject
from docx.text.paragraph import Paragraph as ParagraphObject
from typing import Dict, List, Any, Union  # Union might not be needed here anymore

# To inspect element types
from docx.oxml.ns import qn
import os  # For basename in __main__


def process_docx_file(file_path: str) -> Dict[str, Any]:
    """
    Processes a DOCX file to extract text with table placeholders and a list of table objects.

    Args:
        file_path (str): The path to the .docx file.

    Returns:
        Dict[str, Any]: A dictionary containing:
                        - "text_with_placeholders": Extracted text with table placeholders.
                        - "tables_data": A list of table objects, each with id, position,
                                         caption (None for now), headers, and data.
                        - "source_basename": Basename of the source file.
    """
    try:
        document: DocObject = Document(file_path)

        text_parts: List[str] = []
        # NEW: tables_data will be a list of table detail dictionaries
        tables_data: List[Dict[str, Any]] = []
        table_count = 0

        for block_element in document.element.body:
            if block_element.tag == qn('w:p'):  # Paragraph
                para = ParagraphObject(block_element, document)
                text_parts.append(para.text)
            elif block_element.tag == qn('w:tbl'):  # Table
                table_count += 1
                table_id = f"table{table_count:03d}"

                text_parts.append(f"\n[[INSERT_TABLE:{table_id}]]\n")

                if table_count <= len(document.tables):
                    table_object: TableObject = document.tables[table_count - 1]

                    current_table_rows: List[List[str]] = []
                    for row in table_object.rows:
                        row_data = [cell.text for cell in row.cells]
                        current_table_rows.append(row_data)

                    table_headers: List[str] = []
                    table_actual_data: List[List[str]] = []

                    if current_table_rows:
                        table_headers = current_table_rows[0]  # First row as headers
                        table_actual_data = current_table_rows[1:]  # Remaining rows as data

                    # For now, caption is None. Future enhancement could try to detect it.
                    table_detail = {
                        "id": table_id,
                        "position": table_count,  # 1-based position
                        "caption": None,
                        "headers": table_headers,
                        "data": table_actual_data
                    }
                    tables_data.append(table_detail)
                else:
                    print(f"Warning: Mismatch found for {table_id} in {os.path.basename(file_path)}")

        text_with_placeholders = "\n".join(text_parts)
        text_with_placeholders = text_with_placeholders.replace('\n\n\n', '\n\n').strip()

        return {
            "text_with_placeholders": text_with_placeholders,
            "tables_data": tables_data,  # New key and structure
            "source_basename": os.path.basename(file_path)  # Added for use in main.py
        }

    except Exception as e:
        error_message = f"Error processing DOCX file {file_path}: {str(e)}"
        print(error_message)
        raise ValueError(error_message)


if __name__ == '__main__':
    # ... (The __main__ block can remain similar for direct testing,
    # but its print statements would need to be updated to reflect the new tables_data structure)
    # For brevity, I'll omit the __main__ update here, but you'd adjust it to loop through tables_data list
    # and print each table object's details.
    from docx import Document as CreateDoc

    test_doc_path = "test_schema_output.docx"
    test_doc = CreateDoc()
    test_doc.add_paragraph("This is paragraph one.")
    test_doc.add_paragraph("Text before table one.")

    t1_obj = test_doc.add_table(rows=2, cols=2)
    t1_obj.cell(0, 0).text = "H1"
    t1_obj.cell(0, 1).text = "H2"
    t1_obj.cell(1, 0).text = "D1"
    t1_obj.cell(1, 1).text = "D2"

    test_doc.add_paragraph("Text after table one.")
    test_doc.save(test_doc_path)

    try:
        result = process_docx_file(test_doc_path)
        print("\n--- Extracted Text with Placeholders ---")
        print(result["text_with_placeholders"])
        print("\n--- Tables Data (List of Table Objects) ---")
        for table_obj in result["tables_data"]:
            print(f"\nTable ID: {table_obj['id']}")
            print(f"  Position: {table_obj['position']}")
            print(f"  Caption: {table_obj['caption']}")
            print(f"  Headers: {table_obj['headers']}")
            print(f"  Data Rows: {len(table_obj['data'])}")
            for row in table_obj['data']:
                print(f"    {row}")
        if os.path.exists(test_doc_path):
            os.remove(test_doc_path)  # Clean up test file
    except Exception as e:
        print(f"Error in example usage: {e}")