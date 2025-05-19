import pytest
from fastapi.testclient import TestClient
import os  # For creating dummy files and paths
import sys
from docx import Document as CreateDoc  # To create dummy docx files for testing
from datetime import datetime  # For checking date format, if needed

# Adjust the import path for your FastAPI app instance
# This assumes your FastAPI app instance is named 'app' in 'app/main.py'
# and that pytest is run from the root of your project.
try:
    from app.main import app
except ModuleNotFoundError:
    # Add project root to Python path if running tests from a different working directory
    # or if the app module isn't directly discoverable.
    sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
    from app.main import app


@pytest.fixture(scope="module")
def client():
    """
    Create a TestClient instance for the FastAPI app.
    This fixture will be shared across all tests in this module.
    """
    with TestClient(app) as c:
        yield c


# --- Directory for API test specific fixtures ---
# This helps keep fixtures created by API tests separate from unit test fixtures if needed.
TEST_API_FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "api_test_fixtures")
if not os.path.exists(TEST_API_FIXTURES_DIR):
    os.makedirs(TEST_API_FIXTURES_DIR)


# --- Fixture for creating a sample DOCX file for API tests (schema focused) ---
@pytest.fixture
def sample_api_docx_for_schema_tests():
    """Creates a DOCX file with text and tables for API schema testing."""
    file_path = os.path.join(TEST_API_FIXTURES_DIR, "api_sample_for_schema.docx")
    doc = CreateDoc()
    doc.add_paragraph("API Schema Test: Para 1.")

    table1 = doc.add_table(rows=2, cols=2)  # Table 1
    table1.cell(0, 0).text = "Name"
    table1.cell(0, 1).text = "Age"
    table1.cell(1, 0).text = "Alice"
    table1.cell(1, 1).text = "30"  # Changed to string for consistency with schema items type

    doc.add_paragraph("API Schema Test: Para 2, between tables.")

    table2 = doc.add_table(rows=1, cols=1)  # Table 2 (effectively only a header row)
    table2.cell(0, 0).text = "Note"

    doc.add_paragraph("API Schema Test: Para 3.")
    doc.save(file_path)
    yield file_path  # Provide the path to the test
    # Teardown: remove the file after the test
    if os.path.exists(file_path):
        os.remove(file_path)


# --- Basic API Endpoint Tests ---

def test_read_root(client: TestClient):
    """Test the root endpoint that serves the HTML test page."""
    response = client.get("/")
    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]
    assert b"<title>Document Processing API Test</title>" in response.content


def test_get_status(client: TestClient):
    """Test the health check /status endpoint."""
    response = client.get("/status")
    assert response.status_code == 200
    json_response = response.json()
    assert json_response == {"status": "ok", "message": "API is running"}


def test_custom_api_docs(client: TestClient):
    """Test the custom API documentation placeholder page."""
    response = client.get("/documentation")
    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]
    assert b"<title>API Documentation</title>" in response.content
    assert b"For live, interactive API documentation (Swagger UI)" in response.content


def test_fastapi_auto_docs(client: TestClient):
    """Test that FastAPI's automatic /docs endpoint is available."""
    response = client.get("/docs")
    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]


def test_fastapi_auto_redoc(client: TestClient):
    """Test that FastAPI's automatic /redoc endpoint is available."""
    response = client.get("/redoc")
    assert response.status_code == 200
    assert "text/html" in response.headers["content-type"]


# --- Tests for /process/document/ endpoint ---

def test_process_document_no_file(client: TestClient):
    """Test the /process/document/ endpoint without providing a file."""
    response = client.post("/process/document/")  # No files sent
    assert response.status_code == 422
    json_response = response.json()

    assert "detail" in json_response
    found_file_error = False
    for error_detail_item in json_response["detail"]:
        is_missing_type = error_detail_item.get("type") == "missing"
        loc_info = error_detail_item.get("loc", [])
        # For File(...), loc is often ['body', 'file'] or just ['file']
        is_file_location = "file" in loc_info or ("body" in loc_info and "file" in loc_info)

        if is_missing_type and is_file_location:
            found_file_error = True
            break

    assert found_file_error, f"Error detail should indicate 'file' field is missing. Actual detail: {json_response['detail']}"


def test_process_document_unsupported_file_type(client: TestClient):
    """
    Test the /process/document/ endpoint with an unsupported file type (e.g., .txt).
    It should return a 400 error.
    """
    dummy_file_content = b"This is a dummy text file content for unsupported type test."
    dummy_file_name = "test_dummy_unsupported.txt"

    files = {"file": (dummy_file_name, dummy_file_content, "text/plain")}

    response = client.post("/process/document/?output_format=text", files=files)

    assert response.status_code == 400  # Expecting Bad Request
    json_response = response.json()
    assert "detail" in json_response
    assert "Unsupported file type: 'txt'" in json_response["detail"]
    assert "Supported: docx" in json_response["detail"]

def test_process_document_docx_success_text_output(client: TestClient, sample_api_docx_for_schema_tests: str):
    """
    Test successful processing of a DOCX file through the API, requesting plain text output,
    checking for placeholders and appended table data, and new top-level fields.
    """
    with open(sample_api_docx_for_schema_tests, "rb") as f:
        files = {"file": (os.path.basename(sample_api_docx_for_schema_tests), f,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = client.post("/process/document/?output_format=text", files=files)

    assert response.status_code == 200, f"API Error: {response.text}"
    json_response = response.json()  # API returns JSON even for text output type

    assert json_response["filename"] == os.path.basename(sample_api_docx_for_schema_tests)
    assert json_response["format"] == "text"
    assert "extraction_date" in json_response  # Check new top-level field
    # Quick check for ISO format (doesn't validate fully but checks basic structure)
    assert "T" in json_response["extraction_date"] and "+00:00" in json_response["extraction_date"]
    assert json_response["source_type"] == "docx"  # Check new top-level field

    content = json_response["content"]  # This is the plain text string from format_to_plain_text
    assert "API Schema Test: Para 1." in content
    assert "[[INSERT_TABLE:table001]]" in content
    assert "API Schema Test: Para 2, between tables." in content
    assert "[[INSERT_TABLE:table002]]" in content
    assert "API Schema Test: Para 3." in content

    # Check that the appended table data is present (as per format_to_plain_text.py)
    assert "--- Referenced Table Data ---" in content
    assert "--- table001 ---" in content
    assert "Name\tAge" in content
    assert "Alice\t30" in content
    assert "--- table002 ---" in content
    assert "Note" in content  # Table 2 only had a header row


def test_process_document_docx_success_json_output(client: TestClient, sample_api_docx_for_schema_tests: str):
    """
    Test successful processing of a DOCX file through the API, requesting JSON output,
    checking for the new schema.
    """
    with open(sample_api_docx_for_schema_tests, "rb") as f:
        files = {"file": (os.path.basename(sample_api_docx_for_schema_tests), f,
                          "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        response = client.post("/process/document/?output_format=json", files=files)

    assert response.status_code == 200, f"API Error: {response.text}"
    json_response = response.json()

    # Assert top-level schema fields
    assert json_response["filename"] == os.path.basename(sample_api_docx_for_schema_tests)
    assert json_response["format"] == "json"
    assert "extraction_date" in json_response
    # A more robust date check could be:
    try:
        datetime.fromisoformat(json_response["extraction_date"].replace("Z", "+00:00"))  # Python <3.11 needs Z replaced
    except ValueError:
        try:
            datetime.fromisoformat(json_response["extraction_date"])  # Python 3.11+ handles Z
        except ValueError:
            pytest.fail(f"extraction_date is not a valid ISO 8601 format: {json_response['extraction_date']}")

    assert json_response["source_type"] == "docx"

    # Assert 'content' object structure
    assert "content" in json_response
    content = json_response["content"]

    assert "extracted_text_with_placeholders" in content
    text_with_placeholders = content["extracted_text_with_placeholders"]
    assert "API Schema Test: Para 1." in text_with_placeholders
    assert "[[INSERT_TABLE:table001]]" in text_with_placeholders
    assert "API Schema Test: Para 2, between tables." in text_with_placeholders
    assert "[[INSERT_TABLE:table002]]" in text_with_placeholders
    assert "API Schema Test: Para 3." in text_with_placeholders

    assert "tables" in content
    tables_array = content["tables"]
    assert isinstance(tables_array, list)
    assert len(tables_array) == 2

    # Check table001
    table1_data_from_api = tables_array[0]
    assert table1_data_from_api["id"] == "table001"
    assert table1_data_from_api["position"] == 1
    assert table1_data_from_api["caption"] is None
    assert table1_data_from_api["headers"] == ["Name", "Age"]
    assert table1_data_from_api["data"] == [["Alice", "30"]]  # Age is "30" (string)

    # Check table002
    table2_data_from_api = tables_array[1]
    assert table2_data_from_api["id"] == "table002"
    assert table2_data_from_api["position"] == 2
    assert table2_data_from_api["caption"] is None
    assert table2_data_from_api["headers"] == ["Note"]
    assert table2_data_from_api["data"] == []  # No data rows as it was a single-row table treated as header