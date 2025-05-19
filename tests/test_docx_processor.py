import pytest
import os
from docx import Document as CreateDoc

from app.processing.docx_processor import process_docx_file

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
if not os.path.exists(FIXTURES_DIR):
    os.makedirs(FIXTURES_DIR)


@pytest.fixture
def sample_docx_for_schema_tests():  # Renamed for clarity
    file_path = os.path.join(FIXTURES_DIR, "test_docx_for_schema.docx")
    doc = CreateDoc()
    doc.add_paragraph("Paragraph one.")
    doc.add_paragraph("Text before table 1.")

    t1 = doc.add_table(rows=2, cols=2)
    t1.cell(0, 0).text = "H1_T1"
    t1.cell(0, 1).text = "H2_T1"
    t1.cell(1, 0).text = "D1_T1"
    t1.cell(1, 1).text = "D2_T1"

    doc.add_paragraph("Text between tables.")

    t2 = doc.add_table(rows=1, cols=1)  # Simple table
    t2.cell(0, 0).text = "SingleCell_T2"

    doc.add_paragraph("Text after table 2.")
    doc.save(file_path)
    yield file_path
    if os.path.exists(file_path):
        os.remove(file_path)


@pytest.fixture
def sample_docx_text_only_for_schema():  # Renamed for clarity
    file_path = os.path.join(FIXTURES_DIR, "test_docx_text_only_schema.docx")
    doc = CreateDoc()
    doc.add_paragraph("Hello world for schema test.")
    doc.add_paragraph("This document has no tables.")
    doc.save(file_path)
    yield file_path
    if os.path.exists(file_path):
        os.remove(file_path)


def test_process_docx_text_only_schema(sample_docx_text_only_for_schema):
    result = process_docx_file(sample_docx_text_only_for_schema)

    assert "text_with_placeholders" in result
    assert "tables_data" in result
    assert "Hello world for schema test." in result["text_with_placeholders"]
    assert "[[INSERT_TABLE:" not in result["text_with_placeholders"]
    assert isinstance(result["tables_data"], list)
    assert len(result["tables_data"]) == 0
    assert result["source_basename"] == "test_docx_text_only_schema.docx"


def test_process_docx_with_tables_schema(sample_docx_for_schema_tests):
    result = process_docx_file(sample_docx_for_schema_tests)

    assert "text_with_placeholders" in result
    assert "tables_data" in result

    text = result["text_with_placeholders"]
    tables = result["tables_data"]  # This is now a list of table objects

    assert "Paragraph one." in text
    assert "[[INSERT_TABLE:table001]]" in text
    assert "Text between tables." in text
    assert "[[INSERT_TABLE:table002]]" in text
    assert "Text after table 2." in text

    assert isinstance(tables, list)
    assert len(tables) == 2

    # Check table001 structure
    table1 = tables[0]
    assert table1["id"] == "table001"
    assert table1["position"] == 1
    assert table1["caption"] is None
    assert table1["headers"] == ["H1_T1", "H2_T1"]
    assert table1["data"] == [["D1_T1", "D2_T1"]]

    # Check table002 structure
    table2 = tables[1]
    assert table2["id"] == "table002"
    assert table2["position"] == 2
    assert table2["caption"] is None
    assert table2["headers"] == ["SingleCell_T2"]  # Single cell table, header is the first (only) row
    assert table2["data"] == []  # No data rows if first row is header and only one row exists

    assert result["source_basename"] == "test_docx_for_schema.docx"


def test_process_non_existent_file_schema():  # Renamed for clarity
    with pytest.raises(ValueError) as excinfo:
        process_docx_file("non_existent_document.docx")
    actual_error_message = str(excinfo.value)
    assert "Error processing DOCX file" in actual_error_message
    assert "non_existent_document.docx" in actual_error_message
    assert "Package not found" in actual_error_message